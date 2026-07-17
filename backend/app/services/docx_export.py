"""Render an ADR/KT Markdown document to a .docx file (bytes).

A focused Markdown subset is supported — the shape our templates produce: YAML front-matter
(stripped), ATX headings, the metadata table, bullet lists, fenced code blocks, blockquotes,
and inline **bold** / `code`. Good enough to hand a polished Word doc to reviewers; not a
general Markdown engine.
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_BRAND = RGBColor(0xCC, 0x00, 0x00)
_FRONTMATTER = re.compile(r"^---\n.*?\n---\n", re.DOTALL)
_BOLD = re.compile(r"\*\*(.+?)\*\*")
_CODE = re.compile(r"`([^`]+)`")


def _add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, honoring **bold** and `code` inline spans."""
    # Split on bold/code while keeping delimiters.
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(tok)


def _flush_table(doc: Document, rows: list[list[str]]) -> None:
    rows = [r for r in rows if not all(set(c.strip()) <= {"-", ":"} for c in r)]  # drop |---| separators
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Light Grid Accent 1"
    for r in rows:
        cells = table.add_row().cells
        for i in range(cols):
            cells[i].text = r[i].strip() if i < len(r) else ""


def markdown_to_docx_bytes(markdown: str, title: str) -> bytes:
    markdown = _FRONTMATTER.sub("", markdown, count=1)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            _flush_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # tables (contiguous | ... | rows)
        if line.strip().startswith("|") and line.strip().endswith("|"):
            cells = [c for c in line.strip().strip("|").split("|")]
            table_buf.append(cells)
            i += 1
            continue
        else:
            flush_table()

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=0)
            for run in h.runs:
                run.font.color.rgb = _BRAND
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = _BRAND
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.strip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line.strip()[2:])
        elif line.strip().startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(line.strip()[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif line.strip() == "":
            pass
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)
        i += 1

    flush_table()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
